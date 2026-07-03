from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUT_TS = ROOT / "calibrated_sim_pretraining_acp_20260630_121318.docx"
OUT_LATEST = ROOT / "calibrated_sim_pretraining_acp.docx"
FIG = ROOT / "figures"


def set_font(run, size=10, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_section(section, columns=1):
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.63)
    section.right_margin = Inches(0.63)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    col = cols[0] if cols else OxmlElement("w:cols")
    col.set(qn("w:num"), str(columns))
    col.set(qn("w:space"), "360")
    if not cols:
        sect_pr.append(col)


def switch_columns(doc, columns):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(section, columns)
    return section


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for name, size, bold in [
        ("Title", 24, False),
        ("Heading 1", 10, True),
        ("Heading 2", 10, False),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.14)
    run = p.add_run(text)
    set_font(run, 10)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, 10, bold=(level == 1), italic=(level == 2))
    return p


def add_caption(doc, text, table=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = table
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, 8)
    return p


def add_wide_figure(doc, filename, caption, width=7.0):
    switch_columns(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(FIG / filename), width=Inches(width))
    add_caption(doc, caption)
    switch_columns(doc, 2)


def set_cell_text(cell, text, bold=False, size=7.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for edge in ("top", "left", "bottom", "right"):
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            mar.append(node)
        node.set(qn("w:w"), "70")
        node.set(qn("w:type"), "dxa")


def add_results_table(doc):
    switch_columns(doc, 1)
    add_caption(
        doc,
        "TABLE I. Core five-seed results from the audited aggregate JSON "
        "(mean +/- standard deviation, %).",
        table=True,
    )
    rows = [
        ["Training protocol", "A overall", "B overall", "A 128QAM", "B 128QAM", "B 256QAM"],
        ["Calibrated SimOnly", "71.3+/-1.4", "71.3+/-2.4", "52.7+/-14.4", "8.3+/-5.3", "70.8+/-10.5"],
        ["AWGN-only", "44.6+/-1.8", "40.7+/-2.4", "0.0+/-0.0", "0.0+/-0.0", "95.0+/-6.1"],
        ["5% SimPretrain+RealFT", "84.4+/-1.8", "74.9+/-1.2", "66.0+/-3.3", "2.9+/-2.6", "81.2+/-12.3"],
        ["5% Real-scratch", "69.9+/-13.9", "62.9+/-10.6", "51.3+/-16.7", "29.0+/-8.2", "64.4+/-36.2"],
        ["100% SimPretrain+RealFT", "94.9+/-1.3", "76.9+/-0.3", "85.3+/-11.9", "1.0+/-0.7", "100.0+/-0.0"],
        ["100% Real-scratch", "90.3+/-1.4", "74.5+/-1.6", "70.7+/-9.8", "23.5+/-9.7", "99.8+/-0.4"],
    ]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.75)] + [Inches(1.04)] * 5
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.width = widths[j]
            set_cell_text(cell, text, bold=(i == 0))
    switch_columns(doc, 2)


def add_reference(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"[{number}] {text}")
    set_font(run, 8)


doc = Document()
configure_section(doc.sections[0], 1)
style_document(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
run = title.add_run(
    "Calibrated Simulation Pretraining for Label-Efficient Modulation "
    "Recognition in Experimental OFDM-FSO Links"
)
set_font(run, 24)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.paragraph_format.space_after = Pt(8)
set_font(authors.add_run("Anonymous Authors\nAffiliation withheld for review"), 11)

abstract = doc.add_paragraph()
abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abstract.paragraph_format.space_after = Pt(3)
set_font(abstract.add_run("Abstract-"), 9, bold=True)
set_font(
    abstract.add_run(
        "Experimental free-space optical (FSO) links provide too little labeled data "
        "to train robust modulation recognizers for every channel condition and "
        "acquisition session. This paper investigates calibrated simulation as a "
        "data-efficient pretraining prior for six-class OFDM recognition from QPSK "
        "through 256QAM. Calibrated-Sim-v2.2 models additive noise, lognormal "
        "scintillation, phase errors, bad subcarriers, impulsive noise, and burst "
        "outliers. Across five seeds, using 5% of the in-domain training set increases "
        "accuracy from 69.9+/-13.9% for scratch training to 84.4+/-1.8%, reaching "
        "93.4% of full-data scratch performance. Calibrated simulation alone exceeds "
        "an equal-size AWGN-only corpus by 26.8 percentage points. An independent "
        "session reveals an 18.0+/-1.4-point domain gap dominated by 128QAM collapse. "
        "The results support calibrated pretraining for label efficiency while showing "
        "that aggregate high-order accuracy can conceal class-specific sensitivity."
    ),
    9,
)

keywords = doc.add_paragraph()
keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
keywords.paragraph_format.space_after = Pt(5)
set_font(keywords.add_run("Keywords-"), 9, bold=True)
set_font(
    keywords.add_run(
        "free-space optical communication, OFDM, automatic modulation recognition, "
        "calibrated simulation, sim-to-real transfer, high-order QAM"
    ),
    9,
)

switch_columns(doc, 2)

add_heading(doc, "I. INTRODUCTION")
add_body(
    doc,
    "Automatic modulation recognition allows an adaptive receiver to infer the "
    "transmitted format without explicit signaling. OFDM recognition is difficult "
    "because oscillator offsets, timing errors, channel distortion, and noise alter "
    "both intra-symbol and inter-symbol structure. Classical likelihood- and "
    "feature-based classifiers require accurate models or selected statistics [1], "
    "[2]. Deep models reduce this dependence by learning from IQ samples or derived "
    "representations [3], but their performance remains coupled to the training distribution.",
    indent=False,
)
add_body(
    doc,
    "Recent OFDM recognizers address uncertainty with blind FFT-window estimation "
    "[4], residual learning and radio measurements [5], [6], bi-stream attention [7], "
    "explicit IQ-imbalance and carrier-offset compensation [8], and self-supervised "
    "augmentation [9]. Synthetic-to-OTA evaluation further shows that impairment "
    "diversity and independent-domain testing materially affect conclusions [10].",
)
add_body(
    doc,
    "Optical modulation-format identification often exploits constellation or "
    "amplitude-distribution features [11]. In FSO links, simulation-trained neural "
    "systems can transfer to experimentally distorted optical modes [12]. However, "
    "the label-efficiency of calibrated simulation has not been isolated for "
    "experimental OFDM-FSO modulation recognition under turbulence and session shift.",
)
add_body(
    doc,
    "This paper contributes a six-family calibrated simulator and a controlled "
    "five-seed study. With 5% real data, pretraining improves A-test accuracy by "
    "14.6 points over scratch training; calibrated SimOnly exceeds AWGN-only by "
    "26.8 points. Independent-session evaluation also exposes an 18.0+/-1.4-point "
    "gap dominated by 128QAM rather than by high-order QAM uniformly.",
)

add_heading(doc, "II. EXPERIMENTAL DATASET AND SIGNAL REPRESENTATION")
add_heading(doc, "A. Independent Measurement Sessions", level=2)
add_body(
    doc,
    "The corpus contains QPSK, 16QAM, 32QAM, 64QAM, 128QAM, and 256QAM. "
    "Each frame contains 128 OFDM symbols over 123 active subcarriers. Dataset-A "
    "contains 900 frames split at file level into 648 training, 72 validation, and "
    "180 test frames. Dataset-B is a separate 720-frame session; 144 frames form a "
    "fine-tuning pool and 576 remain for testing. Weak and strong turbulence are "
    "merged for training and retained for stratified evaluation.",
    indent=False,
)
add_body(
    doc,
    "Dataset-B has higher mean PAPR and amplitude variation, with smaller shifts in "
    "phase and IQ statistics. These summaries do not identify one physical cause, "
    "but they establish that B-test is not an exchangeable random split of Dataset-A.",
)
add_wide_figure(
    doc,
    "fig1_ab_distribution.png",
    "Fig. 1. Measured channel statistics for Dataset-A and Dataset-B under weak and "
    "strong turbulence. PAPR and amplitude statistics show a visible session shift.",
    7.0,
)

add_heading(doc, "B. Representation and Recognition Backbone", level=2)
add_body(
    doc,
    "Each frame is centered and RMS-normalized. A 64x64 constellation-density map "
    "(CDM) represents the IQ distribution. A second vector concatenates 16 blind "
    "statistics with a 32-bin radial-amplitude histogram. Four CDM convolution stages "
    "produce 64 features, while a statistics MLP maps 48 inputs to 32 features. Their "
    "96-dimensional concatenation feeds a 128-unit hidden layer and a six-way classifier.",
    indent=False,
)

add_heading(doc, "III. CALIBRATED SIMULATION PRETRAINING")
add_heading(doc, "A. Calibrated-Sim-v2.2", level=2)
add_body(
    doc,
    "The simulator uses the experimental OFDM grid: a 256-point FFT, cyclic prefix "
    "length 16, 123 active subcarriers, and 128 symbols per frame. It applies AWGN, "
    "lognormal scintillation, common and sample-level phase errors, sparse bad "
    "subcarrier or symbol corruption, impulsive noise, and short burst outliers. "
    "SNR ranges and corruption rates depend on modulation order and turbulence. "
    "Calibration targets measured PAPR, amplitude, phase, and outlier statistics, "
    "and parameters are fixed before classifier training. Both calibrated and "
    "AWGN-only corpora contain 12,000 balanced frames.",
    indent=False,
)
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(eq.add_run("x~_k = a_k (x_k + n_k) exp[j(phi + epsilon_k)]"), 10, italic=True)
eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(eq2.add_run("r_k = x~_k + s_k + i_k + b_k"), 10, italic=True)

add_heading(doc, "B. Training Protocols", level=2)
add_body(
    doc,
    "SimOnly trains on calibrated simulation without real adaptation. Real-scratch "
    "trains from random initialization on 5%, 10%, 20%, or 100% of A-train. "
    "SimPretrain+RealFT fine-tunes the SimOnly checkpoint on matched real subsets. "
    "AWGN-only repeats SimOnly with the AWGN corpus. B-FT adapts the best A-trained "
    "model with 36 or 72 Dataset-B frames.",
    indent=False,
)
add_body(
    doc,
    "Pretraining and scratch training use at most 50 epochs at learning rate 1e-3; "
    "batch sizes are 32 and 16. Fine-tuning uses at most 20 epochs, batch size 16, "
    "and learning rate 1e-4. Early stopping selects checkpoints. Every configuration "
    "uses seeds 2026-2030, and results are mean +/- standard deviation over five runs.",
)

add_heading(doc, "IV. EXPERIMENTAL RESULTS AND DISCUSSION")
add_heading(doc, "A. Real-Data Efficiency", level=2)
add_body(
    doc,
    "With 5% of A-train, SimPretrain+RealFT reaches 84.4+/-1.8% on A-test, "
    "compared with 69.9+/-13.9% for Real-scratch. The 14.6-point gain also has "
    "lower seed variability. On B-test, the same comparison is 74.9+/-1.2% versus "
    "62.9+/-10.6%, a 12.0-point gain. The 5% pretrained model reaches 93.4% of "
    "the 90.3+/-1.4% accuracy obtained by scratch training on all A-train frames.",
    indent=False,
)
add_wide_figure(
    doc,
    "fig2_s3_ratio_curves.png",
    "Fig. 2. Accuracy versus Dataset-A training fraction for pretrained fine-tuning "
    "and scratch training. Error bars are standard deviations over five seeds.",
    7.0,
)
add_results_table(doc)

add_heading(doc, "B. Why Calibration Matters", level=2)
add_body(
    doc,
    "Calibrated SimOnly obtains 71.3+/-1.4% on A-test and 71.3+/-2.4% on "
    "B-test. AWGN-only reaches 44.6+/-1.8% and 40.7+/-2.4%. The gains are "
    "26.8 and 30.6 points. Calibration also raises A-test 128QAM from 0.0+/-0.0% "
    "to 52.7+/-14.4%. AWGN-only nevertheless attains high 256QAM accuracy while "
    "failing on 64QAM and 128QAM, indicating a class-specific shortcut rather than "
    "balanced recognition.",
    indent=False,
)
add_wide_figure(
    doc,
    "fig3_awgn_ablation.png",
    "Fig. 3. Equal-size calibrated and AWGN-only simulation ablation. Calibration "
    "improves overall and high-order accuracy and avoids complete 128QAM failure.",
    7.0,
)

add_heading(doc, "C. Cross-Session Failure Analysis", level=2)
add_body(
    doc,
    "The best A-trained model reaches 94.9+/-1.3% on A-test but 76.9+/-0.3% "
    "on B-test, an 18.0+/-1.4-point gap. Its A-test 128QAM and 256QAM accuracies "
    "are 85.3+/-11.9% and 92.7+/-3.9%; on B-test they are 1.0+/-0.7% and "
    "100.0+/-0.0%. Therefore, high-order QAM is not a homogeneous difficulty.",
    indent=False,
)
add_body(
    doc,
    "B-domain fine-tuning provides only a modest correction: 77.2+/-0.7% with "
    "36 frames and 78.2+/-0.2% with 72 frames. The corresponding 128QAM "
    "accuracies remain 1.3+/-1.0% and 1.7+/-1.3%. Calibrated simulation improves "
    "label efficiency in aggregate, but targeted session adaptation is still needed.",
)
add_wide_figure(
    doc,
    "fig4_confusion_matrices.png",
    "Fig. 4. Representative audited confusion matrices for seed 2026. B-test maps "
    "most 128QAM frames to 64QAM or 256QAM, and limited B fine-tuning does not "
    "restore the 128QAM diagonal.",
    6.4,
)

add_heading(doc, "D. Limitations", level=2)
add_body(
    doc,
    "The evidence comes from one optical setup and two sessions, so it does not "
    "establish generalization across sites, weather regimes, or hardware. Calibration "
    "may retain setup-specific priors. The classifier uses preprocessed subcarrier "
    "representations rather than a fully blind waveform, and the B-domain pool cannot "
    "determine whether 128QAM needs more labels, another representation, or explicit "
    "domain alignment.",
    indent=False,
)

add_heading(doc, "V. CONCLUSION")
add_body(
    doc,
    "Calibrated multi-impairment simulation is an effective pretraining source for "
    "experimental OFDM-FSO modulation recognition. With 5% real-data fine-tuning, "
    "it improves in-domain accuracy by 14.6 points over scratch training, while "
    "calibrated SimOnly exceeds AWGN-only by 26.8 points. Independent-session "
    "testing exposes the boundary: the remaining domain gap is dominated by 128QAM. "
    "Future work should calibrate across more sessions and develop class-aware adaptation.",
    indent=False,
)

add_heading(doc, "REFERENCES")
references = [
    "O. A. Dobre, A. Abdi, Y. Bar-Ness, and W. Su, \"Survey of automatic modulation classification techniques: Classical approaches and new trends,\" IET Commun., vol. 1, no. 2, pp. 137-156, 2007, doi: 10.1049/iet-com:20050176.",
    "A. Kumar, S. Majhi, G. Gui, H.-C. Wu, and C. Yuen, \"A survey of blind modulation classification techniques for OFDM signals,\" Sensors, vol. 22, no. 3, p. 1020, 2022, doi: 10.3390/s22031020.",
    "T. J. O'Shea and J. Hoydis, \"An introduction to deep learning for the physical layer,\" IEEE Trans. Cogn. Commun. Netw., vol. 3, no. 4, pp. 563-575, 2017, doi: 10.1109/TCCN.2017.2758370.",
    "M. C. Park and D. S. Han, \"Deep learning-based automatic modulation classification with blind OFDM parameter estimation,\" IEEE Access, vol. 9, pp. 108305-108317, 2021, doi: 10.1109/ACCESS.2021.3102223.",
    "L. Zhang, C. Lin, W. Yan, Q. Ling, and Y. Wang, \"Real-time OFDM signal modulation classification based on deep learning and software-defined radio,\" IEEE Commun. Lett., vol. 25, no. 9, pp. 2988-2992, 2021, doi: 10.1109/LCOMM.2021.3093451.",
    "A. Kumar, K. K. Srinivas, and S. Majhi, \"Automatic modulation classification for adaptive OFDM systems using convolutional neural networks with residual learning,\" IEEE Access, vol. 11, pp. 61013-61024, 2023, doi: 10.1109/ACCESS.2023.3286939.",
    "A. Kumar, M. S. Chaudhari, and S. Majhi, \"Automatic modulation classification for OFDM systems using bi-stream and attention-based CNN-LSTM model,\" IEEE Commun. Lett., vol. 28, no. 3, pp. 552-556, 2024, doi: 10.1109/LCOMM.2023.3348512.",
    "B. Ren, K. C. Teh, H. An, and E. Gunawan, \"OFDM modulation classification using Cross-SKNet with blind IQ imbalance and carrier frequency offset compensation,\" IEEE Trans. Veh. Technol., vol. 73, no. 6, pp. 8389-8403, 2024, doi: 10.1109/TVT.2024.3356606.",
    "B. Ren et al., \"CPAA: Self-supervised cross-view prediction with automatic augmentation for OFDM modulation classification,\" IEEE Trans. Veh. Technol., vol. 74, no. 7, pp. 10536-10550, 2025, doi: 10.1109/TVT.2025.3543800.",
    "R. Nelega, M.-M. Mezei, Z. A. Polgar, G. Kovacs, and E. Puschita, \"Deep learning-based automatic modulation classification for OFDM signals: From synthetic training to OTA evaluation,\" Sensors, vol. 26, no. 10, p. 2945, 2026, doi: 10.3390/s26102945.",
    "Y. Zhao et al., \"Low complexity OSNR monitoring and modulation format identification based on binarized neural networks,\" J. Lightwave Technol., vol. 38, no. 6, pp. 1314-1322, 2020, doi: 10.1109/JLT.2020.2973232.",
    "S. Lohani, E. M. Knutson, and R. T. Glasser, \"Generative machine learning for robust free-space communication,\" Commun. Phys., vol. 3, p. 177, 2020, doi: 10.1038/s42005-020-00444-9.",
]
for index, reference in enumerate(references, start=1):
    add_reference(doc, index, reference)

for section in doc.sections:
    configure_section(section, int(section._sectPr.xpath("./w:cols")[0].get(qn("w:num"), "1")))

doc.save(OUT_TS)
doc.save(OUT_LATEST)
print(OUT_TS)
print(OUT_LATEST)
